import re
from docx import Document
import unicodedata
import pandas as pd

def clean_text(text):
    if not text:
        return ""
    text = text.replace("\xa0", " ")
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()

MIN_YEAR = 1930

def extract_note_and_interpretation(path):

    def normalize_date(text):
        match = re.match(r"(?:(\d{2})/(\d{2})/(\d{4})|(\d{2})/(\d{4}))",text)
        if match:
            if match.group(3):
                jour,mois,annee = match.group(1),match.group(2),match.group(3)
            else:
                jour,mois,annee = "01",match.group(4),match.group(5)
            return f"{jour}/{mois}/{annee}"
        return text

    doc = Document(path)
    full_text = "\n".join([clean_text(p.text) for p in doc.paragraphs])

    sexe = ""
    date_naissance = ""
    charges_virales = []
    t_cd4 = []
    resultats = ""

    # === 📊 Lecture du tableau (ligne par ligne)
    for table in doc.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            if len(cells) < 3:
                continue

            # Ligne contenant date de naissance et sexe
            if not date_naissance and re.match(r"\d{2}/\d{2}/\d{4}", cells[0]):
                try:
                    parsed_naissance = pd.to_datetime(cells[0], dayfirst=True, errors="coerce")
                    if parsed_naissance and parsed_naissance.year >= MIN_YEAR:
                        date_naissance = parsed_naissance.strftime("%d/%m/%Y")
                except:
                    pass
                if "f" in cells[1].lower():
                    sexe = "Féminin"
                elif "m" in cells[1].lower():
                    sexe = "Masculin"
                else:
                    sexe = "Autre"
                continue

            # 2️⃣ Extraction des virémies (dates en col 2, valeurs en col 5)
            if len(cells) >= 6 and re.match(r"(?:(\d{2})/(\d{2})/(\d{4})|(\d{2})/(\d{4}))", cells[2]):
                try:
                    date_cell = cells[2].strip()
                    raw_date = normalize_date(date_cell)
                    if raw_date:
                        formated_date =  pd.to_datetime(raw_date,dayfirst=False, errors="coerce")
                        viremie = re.sub(r"\D", "", cells[5])
                        if viremie:
                            charges_virales.append({
                                "valeur": viremie,
                                "date":formated_date.strftime("%Y-%m-%d")

                            })
                except:
                    continue

            # 2️⃣ Extraction des cd4 (dates en col 2, valeurs en col 3)
            if len(cells) >= 4 and re.match(r"\b(?:(\d{2})/(\d{2})/(\d{4})|(\d{2})/(\d{4}))\b", cells[2]):
                try:
                    date_cell = cells[2].strip()
                    raw_date = normalize_date(date_cell)
                    if raw_date:
                            formated_date =  pd.to_datetime(raw_date,dayfirst=False, errors="coerce")
                            cd4 = re.sub(r"\D", "", cells[3])
                            if cd4:
                                t_cd4.append({
                                    "valeur": cd4,
                                    "date": formated_date.strftime("%Y-%m-%d")
                                })
                except:
                    continue

    # --- 🧪 Extraction section RESULTATS ---
    res_match = re.search(r"RESULTATS\s*:?\s*(.*?)\s*(?=NOTE|INTERPRETATION|\Z)", full_text, re.DOTALL | re.IGNORECASE)
    if res_match:
        resultats = res_match.group(1).strip()

    # --- 📝 Extraction NOTE ---
    note_text = ""
    note_pattern = r"NOTE\s*:?\s*(.*?)(?=INTERPRÉTATION\s+VIROLOGIQUE|\Z)"
    note_match = re.search(note_pattern, full_text, flags=re.IGNORECASE | re.DOTALL)
    if note_match:
        note_text = note_match.group(1).strip()
    else:
        note_text = "NOTE : Non trouvée"

    # --- 🔬 Extraction INTERPRÉTATION ---
    interpretation_text = ""
    viro_pattern = r"INTERPRÉTATION\s+VIROLOGIQUE.*?\)?\s*(.*)"
    viro_match = re.search(viro_pattern, full_text, flags=re.IGNORECASE | re.DOTALL)
    if viro_match:
        interpretation_text = viro_match.group(1).strip()
    else:
        interpretation_text = "INTERPRETATION VIROLOGIQUE : Non trouvée"

    return {
        "sexe": sexe,
        "date_naissance": date_naissance,
        "charges_virales": charges_virales,
        "taux_cd4": t_cd4,
        "resultats": resultats,
        "note": note_text,
        "interpretation": interpretation_text
    }